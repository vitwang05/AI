from fastapi import FastAPI, UploadFile, File, HTTPException, Depends, Request, Response
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from pydantic import BaseModel
import os
import json
from app.qa_chain import create_qa_chain, answer_question
from app.document_processor import process_document, setup_pinecone_index, extract_structured_terms
from typing import List, Dict, Optional
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
from datetime import datetime, timedelta
from jose import JWTError, jwt
from passlib.context import CryptContext
from collections import defaultdict

# Cấu hình JWT
SECRET_KEY = "your-secret-key"  # Thay đổi thành một key bảo mật trong môi trường production
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 30

# Cấu hình password hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# Model cho user
class User(BaseModel):
    username: str
    role: str

# Model cho token
class Token(BaseModel):
    access_token: str
    token_type: str

# Model cho user trong database
class UserInDB(User):
    hashed_password: str

# Mock database - trong môi trường production nên sử dụng database thật
fake_users_db = {
    "user": {
        "username": "user",
        "role": "user",
        "hashed_password": pwd_context.hash("user123")
    },
    "admin": {
        "username": "admin",
        "role": "admin",
        "hashed_password": pwd_context.hash("admin123")
    }
}

app = FastAPI()

# Cấu hình CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:3001"],  # Frontend URLs
    allow_credentials=True,
    allow_methods=["*"],  # Cho phép tất cả các HTTP methods
    allow_headers=["*"],  # Cho phép tất cả các headers
)

# OAuth2 scheme
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="login")

# Hàm xác thực user
def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_user(db, username: str):
    if username in db:
        user_dict = db[username]
        return UserInDB(**user_dict)
    return None

def authenticate_user(fake_db, username: str, password: str):
    user = get_user(fake_db, username)
    if not user:
        return False
    if not verify_password(password, user.hashed_password):
        return False
    return user

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(token: str = Depends(oauth2_scheme)):
    credentials_exception = HTTPException(
        status_code=401,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
    user = get_user(fake_users_db, username)
    if user is None:
        raise credentials_exception
    return user

@app.post("/login", response_model=Token)
async def login(form_data: OAuth2PasswordRequestForm = Depends()):
    user = authenticate_user(fake_users_db, form_data.username, form_data.password)
    if not user:
        raise HTTPException(
            status_code=401,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user.username}, expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer"}

def process_json(data, qa):
    results = []
    #Copy file output.json
    document = []
    #Mỗi lần lưu result thì sẽ thêm answer vào document ở mục cuối cùng đang chạy đến trong output.json
    def group_by_program(documents):
        result = defaultdict(list)
        for item in documents:
            result[item.metadata.get('program')].append({   
                "title": item.metadata.get('article_title'), 
                "text": item.page_content})
        return dict(result)
    
    def process_item(item, parent_title=""):
        title = item.get("title", "")
        sentence = f"{parent_title}\n{title}" if parent_title else title
        full_title = f"{parent_title} > {title}" if parent_title else title
        sub_items = item.get("sub_items", [])

        if not sub_items:
            answer, documents = answer_question(full_title, qa)
            grouped_documents = group_by_program(documents)
            results.append({
                "sentence": sentence,
                "question": full_title,
                "answer": answer,
                "documents": grouped_documents
            })
            document.append({
                "title": title,
                "answer": answer,
                "documents": grouped_documents
            })
        else:
            processed_sub_items = []
            for sub in sub_items:
                sub_title = sub.get("title", "")
                sub_sentence = f"{sentence}\n{sub_title}"
                full_sub_title = f"{full_title} > {sub_title}"
                details = sub.get("details", [])
                if not details:
                    answer, documents = answer_question(full_sub_title, qa)
                    grouped_documents = group_by_program(documents)
                    results.append({
                        "sentence": sub_sentence,
                        "question": full_sub_title,
                        "answer": answer,
                        "documents": grouped_documents
                    })
                    processed_sub_items.append({
                        "title": sub_title,
                        "answer": answer,
                        "documents": grouped_documents
                    })
                else:
                    processed_details = []
                    for detail in details:
                        detail_title = detail.get("title", "")
                        detail_sentence = f"{sub_sentence}\n{detail_title}"
                        full_detail_title = f"{full_sub_title} > {detail_title}"
                        sub_details = detail.get("sub_details", [])
                        if not sub_details:
                            answer, documents = answer_question(full_detail_title, qa)
                            grouped_documents = group_by_program(documents)
                            results.append({
                                "sentence": detail_sentence,
                                "question": full_detail_title,
                                "answer": answer,
                                "documents": grouped_documents
                            })
                            processed_details.append({
                                "title": detail_title,
                                "answer": answer,
                                "documents": grouped_documents
                            })
                        else:
                            processed_sub_details = []
                            for sub_detail in sub_details:
                                sub_detail_title = sub_detail.get("title", "")
                                sub_detail_sentence = f"{detail_sentence}\n{sub_detail_title}"
                                full_sub_detail_title = f"{full_detail_title} > {sub_detail_title}"
                                answer, documents = answer_question(full_sub_detail_title, qa)
                                grouped_documents = group_by_program(documents)
                                results.append({
                                    "sentence": sub_detail_sentence,
                                    "question": full_sub_detail_title,
                                    "answer": answer,
                                    "documents": grouped_documents
                                })
                                processed_sub_details.append({
                                    "title": sub_detail_title,
                                    "answer": answer,
                                    "documents": grouped_documents
                                })
                            processed_details.append({
                                "title": detail_title,
                                "sub_details": processed_sub_details
                            })
                    processed_sub_items.append({
                        "title": sub_title,
                        "details": processed_details
                    })
            document.append({
                "title": title,
                "sub_items": processed_sub_items
            })

    # Xử lý từng item trong danh sách
    if isinstance(data, list):
        for item in data:
            process_item(item)
    else:
        process_item(data)
    
    return results, document

@app.post("/uploadVBPL")
async def upload_fileVBPL(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    if current_user.role != "admin":
        raise HTTPException(
            status_code=403,
            detail="Only admin users can upload VBPL files"
        )
    try:
        temp_dir = "VBPL"
        os.makedirs(temp_dir, exist_ok=True)
        file_path = os.path.join(temp_dir, file.filename)
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        return {
            "message": "File uploaded successfully",
            "filename": file.filename,
            "file_path": file_path.replace("\\", "/")
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Lỗi upload file: {str(e)}")

@app.post("/learn")
async def learn_file(
    file_path: str,
    current_user: User = Depends(get_current_user)
):
    if current_user.role != "admin":
        raise HTTPException(
            status_code=403,
            detail="Only admin users can learn files"
        )
    try:
        start_time = time.time()
        # Convert URL-encoded path back to normal path
        file_path = file_path.replace("%2F", "/")
        
        # Kiểm tra file tồn tại
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail=f"File not found: {file_path}")

        # Kiểm tra file có phải là PDF hoặc DOCX không
        if not file_path.lower().endswith(('.pdf', '.docx')):
            raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")
        
        process_document(file_path)
        
        end_time = time.time()
        processing_time = end_time - start_time
        
        return {
            "message": "File processed successfully"
        }
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")
  
@app.post("/uploadVBNB")
async def upload_file(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    try:
        # Create temp directory if it doesn't exist
        temp_dir = "temp"
        os.makedirs(temp_dir, exist_ok=True)
        
        # Save uploaded file
        file_path = os.path.join(temp_dir, file.filename)
        with open(file_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        return {
            "message": "File uploaded successfully",
            "filename": file.filename,
            "file_path": file_path.replace("\\", "/")
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Lỗi upload file: {str(e)}")

@app.post("/process")
async def process_file(
    file_path: str,
    start_page: int,
    end_page: int,
    current_user: User = Depends(get_current_user)
):
    try:
        start_time = time.time()
        # Convert URL-encoded path back to normal path
        file_path = file_path.replace("%2F", "/")
        
        # Kiểm tra file tồn tại
        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail=f"File not found: {file_path}")

        # Kiểm tra file có phải là PDF hoặc DOCX không
        if not file_path.lower().endswith(('.pdf', '.docx')):
            raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")

        # Kiểm tra số trang hợp lệ
        if start_page < 1 or end_page < start_page:
            raise HTTPException(status_code=400, detail="Invalid page range")

        # Gọi hàm xử lý văn bản
        try:
            structured_terms = extract_structured_terms(file_path, start_page, end_page)
            if not structured_terms:
                raise HTTPException(status_code=400, detail="No content found in the specified page range")
        except Exception as e:
            print(f"Error extracting structured terms: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

        # Lưu kết quả trung gian vào thư mục temp
        temp_dir = "json_output"
        os.makedirs(temp_dir, exist_ok=True)
        output_json = os.path.join(temp_dir, "output.json")
        
        with open(output_json, "w", encoding="utf-8") as f:
            json.dump(structured_terms, f, ensure_ascii=False, indent=2)

        # Tạo chuỗi hỏi đáp
        try:
            qa_chain = create_qa_chain()
        except Exception as e:
            print(f"Error creating QA chain: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Error creating QA chain: {str(e)}")

        # Đọc dữ liệu đã trích xuất
        try:
            with open(output_json, "r", encoding="utf-8") as f:
                json_data = json.load(f)
        except Exception as e:
            print(f"Error reading output file: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Error reading output file: {str(e)}")

        # Xử lý dữ liệu JSON
        try:
            results, document = process_json(json_data, qa_chain)
            if not results:
                raise HTTPException(status_code=400, detail="No results generated from the content")
        except Exception as e:
            print(f"Error processing JSON data: {str(e)}")
            print(f"JSON data: {json.dumps(json_data, indent=2)}")
            raise HTTPException(status_code=500, detail=f"Error processing JSON data: {str(e)}")

        # Lưu kết quả cuối cùng vào thư mục output
        output_dir = "output"
        os.makedirs(output_dir, exist_ok=True)

        document_dir = "document"
        os.makedirs(document_dir, exist_ok=True)

        # print(document)
        
        # Lấy tên file gốc
        original_filename = os.path.basename(file_path)
        # Tạo timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        # Tạo tên file kết quả
        document_filename = f"{os.path.splitext(original_filename)[0]}_{timestamp}.json"
        document_json = os.path.join(document_dir, document_filename)
        
        results_filename = f"{os.path.splitext(original_filename)[0]}_{timestamp}.json"
        results_json = os.path.join(output_dir, results_filename)
        # Tính thời gian xử lý
        end_time = time.time()
        processing_time = end_time - start_time

        # Thêm thời gian xử lý vào kết quả
        results.append({"process_time": processing_time})
        
        try:
            with open(results_json, "w", encoding="utf-8") as f:
                json.dump(results, f, ensure_ascii=False, indent=2)
            with open(document_json, "w", encoding="utf-8") as f:
                json.dump(document, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"Error writing results file: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Error writing results file: {str(e)}")
        
        return {
            "results": results,
            "processing_time": processing_time,
            "filename": results_filename
        }
    except HTTPException:
        raise
    except Exception as e:
        print(f"Unexpected error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Unexpected error: {str(e)}")

@app.get("/process-results")
async def get_process_results():
    try:
        output_dir = "output"
        if not os.path.exists(output_dir):
            return []
            
        results = []
        for filename in os.listdir(output_dir):
            if filename.endswith('.json'):
                file_path = os.path.join(output_dir, filename)
                modified_time = os.path.getmtime(file_path)
                results.append({
                    "filename": filename,
                    "modified_time": modified_time
                })
        
        # Sắp xếp theo thời gian sửa đổi mới nhất
        results.sort(key=lambda x: x["modified_time"], reverse=True)
        return results
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/process-results/{filename}")
async def get_process_result(
    filename: str,
    current_user: User = Depends(get_current_user)
):
    try:
        output_dir = "output"
        document_dir = "document"
        
        # Tìm file có chứa tên file gốc trong tên
        matching_files = [f for f in os.listdir(output_dir) if filename in f and f.endswith('.json')]
        matching_document_files = [f for f in os.listdir(document_dir) if filename in f and f.endswith('.json')]
        
        if not matching_files:
            raise HTTPException(status_code=404, detail="Result not found")
        if not matching_document_files:
            raise HTTPException(status_code=404, detail="Document not found")
            
        # Lấy file đầu tiên khớp với tên file
        results_file = os.path.join(output_dir, matching_files[0])
        document_file = os.path.join(document_dir, matching_document_files[0])
        
        try:
            with open(results_file, "r", encoding="utf-8") as f:
                results_data = json.load(f)
                print("Raw results data:", results_data)  # Debug log
        except Exception as e:
            print(f"Error reading results file: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Error reading results file: {str(e)}")
            
        try:
            with open(document_file, "r", encoding="utf-8") as f:
                document_data = json.load(f)
                print("Raw document data:", document_data)  # Debug log
        except Exception as e:
            print(f"Error reading document file: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Error reading document file: {str(e)}")
            
        # Kiểm tra xem results_data có phải là list và có phần tử cuối cùng không
        if not isinstance(results_data, list) or not results_data:
            raise HTTPException(status_code=500, detail="Invalid results data format")
            
        # Kiểm tra xem phần tử cuối cùng có process_time không
        last_result = results_data[-1]
        if not isinstance(last_result, dict) or "process_time" not in last_result:
            raise HTTPException(status_code=500, detail="Process time not found in results")
            
        return {
            "filename": filename,
            "results": results_data,
            "document": document_data,
            "process_time": last_result["process_time"]
        }
    except HTTPException as he:
        raise he
    except Exception as e:
        print(f"Error reading results: {str(e)}")  # Debug log
        raise HTTPException(status_code=500, detail=f"Error reading results: {str(e)}")

@app.post("/generate-docx")
async def generate_docx(request: Request):
    try:
        data = await request.json()
        filename = data.get('filename')
        print(f"Received request for filename: {filename}")
        
        if not filename:
            raise HTTPException(status_code=400, detail="Filename is required")

        # Tìm file trong thư mục document
        document_dir = "document"
        matching_files = [f for f in os.listdir(document_dir) if f.endswith('.json') and filename in f]
        
        if not matching_files:
            print(f"No matching files found for: {filename}")
            raise HTTPException(status_code=404, detail="Document file not found")
            
        # Lấy file đầu tiên tìm thấy
        document_file = matching_files[0]
        json_path = os.path.join(document_dir, document_file)
        print(f"Found matching file: {json_path}")

        # Đọc file JSON
        with open(json_path, 'r', encoding='utf-8') as f:
            document = json.load(f)
        print(f"Successfully loaded JSON data with {len(document)} items")

        # Tạo file DOCX
        doc = Document()
        
        # Thêm tiêu đề
        doc.add_heading('Kết quả phân tích', 0)
        
        # Thêm từng mục
        for item in document:
            # Thêm tiêu đề chính
            doc.add_heading(item['title'], level=1)
            
            # Thêm câu trả lời
            if 'answer' in item:
                p = doc.add_paragraph()
                p.add_run('AI trả lời: ').bold = True
                p.add_run(item['answer'])
            
            # Thêm tài liệu tham khảo
            if 'documents' in item:
                doc.add_heading('Tài liệu tham khảo:', level=2)
                for source, docs in item['documents'].items():
                    doc.add_heading(source, level=3)
                    for doc_item in docs:
                        p = doc.add_paragraph()
                        p.add_run(doc_item['title']).bold = True
                        p.add_run('\n' + doc_item['text'])
            
            # Xử lý sub_items nếu có
            if 'sub_items' in item:
                for sub_item in item['sub_items']:
                    # Thêm tiêu đề phụ
                    doc.add_heading(sub_item['title'], level=2)
                    
                    # Thêm câu trả lời
                    if 'answer' in sub_item:
                        p = doc.add_paragraph()
                        p.add_run('AI trả lời: ').bold = True
                        p.add_run(sub_item['answer'])
                    
                    # Thêm tài liệu tham khảo
                    if 'documents' in sub_item:
                        doc.add_heading('Tài liệu tham khảo:', level=3)
                        for source, docs in sub_item['documents'].items():
                            doc.add_heading(source, level=4)
                            for doc_item in docs:
                                p = doc.add_paragraph()
                                p.add_run(doc_item['title']).bold = True
                                p.add_run('\n' + doc_item['text'])
                    
                    # Xử lý details nếu có
                    if 'details' in sub_item:
                        for detail in sub_item['details']:
                            # Thêm tiêu đề chi tiết
                            doc.add_heading(detail['title'], level=3)
                            
                            # Thêm câu trả lời
                            if 'answer' in detail:
                                p = doc.add_paragraph()
                                p.add_run('AI trả lời: ').bold = True
                                p.add_run(detail['answer'])
                            
                            # Thêm tài liệu tham khảo
                            if 'documents' in detail:
                                doc.add_heading('Tài liệu tham khảo:', level=4)
                                for source, docs in detail['documents'].items():
                                    doc.add_heading(source, level=5)
                                    for doc_item in docs:
                                        p = doc.add_paragraph()
                                        p.add_run(doc_item['title']).bold = True
                                        p.add_run('\n' + doc_item['text'])
                            
                            # Xử lý sub_details nếu có
                            if 'sub_details' in detail:
                                for sub_detail in detail['sub_details']:
                                    # Thêm tiêu đề chi tiết phụ
                                    doc.add_heading(sub_detail['title'], level=4)
                                    
                                    # Thêm câu trả lời
                                    if 'answer' in sub_detail:
                                        p = doc.add_paragraph()
                                        p.add_run('AI trả lời: ').bold = True
                                        p.add_run(sub_detail['answer'])
                                    
                                    # Thêm tài liệu tham khảo
                                    if 'documents' in sub_detail:
                                        doc.add_heading('Tài liệu tham khảo:', level=5)
                                        for source, docs in sub_detail['documents'].items():
                                            doc.add_heading(source, level=6)
                                            for doc_item in docs:
                                                p = doc.add_paragraph()
                                                p.add_run(doc_item['title']).bold = True
                                                p.add_run('\n' + doc_item['text'])

        # Tạo tên file DOCX
        docx_filename = f"{os.path.splitext(filename)[0]}.docx"
        docx_path = os.path.join(document_dir, docx_filename)
        
        # Lưu file DOCX
        doc.save(docx_path)
        print(f"Successfully created DOCX file: {docx_path}")

        # Trả về file DOCX
        return FileResponse(
            path=docx_path,
            filename=docx_filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        print(f"Error generating DOCX: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating DOCX: {str(e)}")

@app.get("/files")
async def get_files(
    directory: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    try:
        if directory == "VBPL" and current_user.role != "admin":
            raise HTTPException(
                status_code=403,
                detail="Only admin users can access VBPL files"
            )

        if directory:
            dir_path = directory
        else:
            dir_path = "temp"

        if not os.path.exists(dir_path):
            return {directory: []}

        files = []
        for filename in os.listdir(dir_path):
            file_path = os.path.join(dir_path, filename)
            if os.path.isfile(file_path):
                stat = os.stat(file_path)
                files.append({
                    "name": filename,
                    "path": file_path.replace("\\", "/"),
                    "size": stat.st_size,
                    "modified": stat.st_mtime
                })

        return {directory: files}
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error getting files: {str(e)}")

@app.delete("/files")
async def delete_file(
    path: str,
    current_user: User = Depends(get_current_user)
):
    try:
        # Convert URL-encoded path back to normal path
        path = path.replace("%2F", "/")
        
        # Kiểm tra quyền xóa file
        if path.startswith("VBPL/") and current_user.role != "admin":
            raise HTTPException(
                status_code=403,
                detail="Only admin users can delete VBPL files"
            )

        if not os.path.exists(path):
            raise HTTPException(status_code=404, detail="File not found")

        os.remove(path)
        return {"message": "File deleted successfully"}
    except HTTPException as he:
        raise he
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting file: {str(e)}")